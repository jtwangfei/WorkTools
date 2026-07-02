from pathlib import Path

import pytest
from docx import Document

from worktools.word_merge import (
    DuplicatePlaceholderError,
    InputDocumentError,
    MergeModeNotImplementedError,
    MissingPlaceholderError,
    OutputExistsError,
    TemplatePlaceholderError,
    WordMergeOptions,
    merge_documents,
    resolve_template_path,
)
from worktools.word_merge_cli import build_parser


def save_docx(path: Path, paragraphs: list[str], style: str | None = None) -> None:
    document = Document()
    for text in paragraphs:
        paragraph = document.add_paragraph(text)
        if style:
            paragraph.style = style
    document.save(path)


def save_template(path: Path, placeholder: str = "{{content}}", style: str = "Quote") -> None:
    document = Document()
    document.add_paragraph("Header")
    paragraph = document.add_paragraph(placeholder)
    paragraph.style = style
    document.add_paragraph("Footer")
    document.save(path)


def test_resolve_template_path_uses_existing_file(tmp_path: Path) -> None:
    template = tmp_path / "custom.docx"
    save_template(template)

    assert resolve_template_path(str(template), tmp_path) == template


def test_resolve_template_path_uses_named_template(tmp_path: Path) -> None:
    templates_dir = tmp_path / "templates"
    templates_dir.mkdir()
    template = templates_dir / "report.docx"
    save_template(template)

    assert resolve_template_path("report", tmp_path) == template


def test_merge_documents_inserts_inputs_in_order_and_uses_template_style(tmp_path: Path) -> None:
    templates_dir = tmp_path / "templates"
    templates_dir.mkdir()
    template = templates_dir / "report.docx"
    first = tmp_path / "first.docx"
    second = tmp_path / "second.docx"
    output = tmp_path / "merged.docx"

    save_template(template, style="Quote")
    save_docx(first, ["First A", "First B"])
    save_docx(second, ["Second A"])

    merge_documents(
        WordMergeOptions(
            template="report",
            inputs=[first, second],
            output=output,
            base_dir=tmp_path,
        )
    )

    document = Document(output)
    assert [paragraph.text for paragraph in document.paragraphs] == [
        "Header",
        "First A",
        "First B",
        "Second A",
        "Footer",
    ]
    assert [paragraph.style.name for paragraph in document.paragraphs[1:4]] == [
        "Quote",
        "Quote",
        "Quote",
    ]


def test_merge_documents_rejects_missing_placeholder(tmp_path: Path) -> None:
    template = tmp_path / "template.docx"
    input_doc = tmp_path / "input.docx"
    document = Document()
    document.add_paragraph("No marker")
    document.save(template)
    save_docx(input_doc, ["Body"])

    with pytest.raises(MissingPlaceholderError):
        merge_documents(
            WordMergeOptions(
                template=str(template),
                inputs=[input_doc],
                output=tmp_path / "out.docx",
            )
        )


def test_merge_documents_rejects_duplicate_placeholders(tmp_path: Path) -> None:
    template = tmp_path / "template.docx"
    input_doc = tmp_path / "input.docx"
    document = Document()
    document.add_paragraph("{{content}}")
    document.add_paragraph("{{content}}")
    document.save(template)
    save_docx(input_doc, ["Body"])

    with pytest.raises(DuplicatePlaceholderError):
        merge_documents(
            WordMergeOptions(
                template=str(template),
                inputs=[input_doc],
                output=tmp_path / "out.docx",
            )
        )


def test_merge_documents_rejects_placeholder_with_other_text(tmp_path: Path) -> None:
    template = tmp_path / "template.docx"
    input_doc = tmp_path / "input.docx"
    save_template(template, placeholder="Before {{content}}")
    save_docx(input_doc, ["Body"])

    with pytest.raises(TemplatePlaceholderError):
        merge_documents(
            WordMergeOptions(
                template=str(template),
                inputs=[input_doc],
                output=tmp_path / "out.docx",
            )
        )


def test_merge_documents_rejects_non_docx_input(tmp_path: Path) -> None:
    template = tmp_path / "template.docx"
    input_doc = tmp_path / "input.txt"
    save_template(template)
    input_doc.write_text("Body", encoding="utf-8")

    with pytest.raises(InputDocumentError):
        merge_documents(
            WordMergeOptions(
                template=str(template),
                inputs=[input_doc],
                output=tmp_path / "out.docx",
            )
        )


def test_merge_documents_rejects_existing_output_without_overwrite(tmp_path: Path) -> None:
    template = tmp_path / "template.docx"
    input_doc = tmp_path / "input.docx"
    output = tmp_path / "out.docx"
    save_template(template)
    save_docx(input_doc, ["Body"])
    save_docx(output, ["Existing"])

    with pytest.raises(OutputExistsError):
        merge_documents(WordMergeOptions(template=str(template), inputs=[input_doc], output=output))


def test_merge_documents_rejects_summarize_mode(tmp_path: Path) -> None:
    template = tmp_path / "template.docx"
    input_doc = tmp_path / "input.docx"
    save_template(template)
    save_docx(input_doc, ["Body"])

    with pytest.raises(MergeModeNotImplementedError):
        merge_documents(
            WordMergeOptions(
                template=str(template),
                inputs=[input_doc],
                output=tmp_path / "out.docx",
                mode="summarize",
            )
        )


def test_cli_parser_builds_word_merge_options() -> None:
    parser = build_parser()

    args = parser.parse_args(
        ["--template", "report", "a.docx", "b.docx", "-o", "merged.docx", "--overwrite"]
    )

    assert args.template == "report"
    assert args.inputs == ["a.docx", "b.docx"]
    assert args.output == "merged.docx"
    assert args.mode == "direct"
    assert args.overwrite is True
