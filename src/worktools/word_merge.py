from __future__ import annotations

from dataclasses import dataclass, field
from pathlib import Path
from typing import Literal

from docx import Document
from docx.document import Document as DocxDocument
from docx.text.paragraph import Paragraph

PLACEHOLDER = "{{content}}"
MergeMode = Literal["direct", "summarize"]


class WordMergeError(Exception):
    """Base error for Word merge failures."""


class TemplateNotFoundError(WordMergeError):
    """Raised when a template cannot be resolved."""


class MissingPlaceholderError(WordMergeError):
    """Raised when the template does not contain the content placeholder."""


class DuplicatePlaceholderError(WordMergeError):
    """Raised when the template contains more than one content placeholder."""


class TemplatePlaceholderError(WordMergeError):
    """Raised when the placeholder paragraph violates the template contract."""


class InputDocumentError(WordMergeError):
    """Raised when an input document is invalid."""


class OutputExistsError(WordMergeError):
    """Raised when the output already exists and overwrite is disabled."""


class MergeModeNotImplementedError(WordMergeError):
    """Raised when a reserved merge mode is requested."""


@dataclass(frozen=True)
class WordMergeOptions:
    template: str
    inputs: list[Path]
    output: Path
    mode: MergeMode = "direct"
    overwrite: bool = False
    base_dir: Path = field(default_factory=Path.cwd)


@dataclass(frozen=True)
class WordMergeResult:
    output: Path


def resolve_template_path(template: str, base_dir: Path | None = None) -> Path:
    base = base_dir or Path.cwd()
    direct_path = Path(template)
    if direct_path.is_file():
        return direct_path

    for suffix in (".docx", ".dotx"):
        candidate = base / "templates" / f"{template}{suffix}"
        if candidate.is_file():
            return candidate

    raise TemplateNotFoundError(f"Template not found: {template}")


def merge_documents(options: WordMergeOptions) -> WordMergeResult:
    if options.mode == "summarize":
        raise MergeModeNotImplementedError("Summarize mode is reserved but not implemented.")
    if options.mode != "direct":
        raise WordMergeError(f"Unsupported merge mode: {options.mode}")
    if options.output.exists() and not options.overwrite:
        raise OutputExistsError(f"Output already exists: {options.output}")

    template_path = resolve_template_path(options.template, options.base_dir)
    _validate_inputs(options.inputs)

    document = Document(template_path)
    placeholder = _find_placeholder(document)
    style_name = placeholder.style.name
    insert_after = placeholder._p

    for input_path in options.inputs:
        source = Document(input_path)
        for source_paragraph in source.paragraphs:
            if source_paragraph.text:
                insert_after = _insert_paragraph_after(
                    document,
                    insert_after,
                    source_paragraph.text,
                    style_name,
                )

    _remove_paragraph(placeholder)
    options.output.parent.mkdir(parents=True, exist_ok=True)
    document.save(options.output)
    return WordMergeResult(output=options.output)


def _validate_inputs(inputs: list[Path]) -> None:
    if not inputs:
        raise InputDocumentError("At least one input document is required.")
    for input_path in inputs:
        if not input_path.is_file():
            raise InputDocumentError(f"Input document not found: {input_path}")
        if input_path.suffix.lower() != ".docx":
            raise InputDocumentError(f"Input document must be a .docx file: {input_path}")


def _find_placeholder(document: DocxDocument) -> Paragraph:
    matches = [paragraph for paragraph in document.paragraphs if PLACEHOLDER in paragraph.text]
    if not matches:
        raise MissingPlaceholderError(f"Template must contain {PLACEHOLDER}.")
    if len(matches) > 1:
        raise DuplicatePlaceholderError(f"Template must contain only one {PLACEHOLDER}.")

    placeholder = matches[0]
    if placeholder.text.strip() != PLACEHOLDER:
        raise TemplatePlaceholderError(f"{PLACEHOLDER} must be the only text in its paragraph.")
    return placeholder


def _remove_paragraph(paragraph: Paragraph) -> None:
    element = paragraph._element
    element.getparent().remove(element)
    paragraph._p = paragraph._element = None


def _insert_paragraph_after(
    document: DocxDocument,
    previous_element,
    text: str,
    style_name: str,
):
    paragraph = document.add_paragraph()
    paragraph.text = text
    paragraph.style = style_name
    previous_element.addnext(paragraph._p)
    return paragraph._p
